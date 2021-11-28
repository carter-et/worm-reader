"""
A minority* of the code seen here was created by jocelynh.
https://github.com/redoctopus/Worm-TFIDF/blob/master/scrape_site.py
"""
import re
import requests
import docx
import math

from bs4 import BeautifulSoup as bs
from docx import Document


FIRST_PAGE = 'https://parahumans.wordpress.com/category/stories-arcs-1-10/arc-1-gestation/1-01/'
TAG_RE = re.compile(r'<[^>]+>')
headers = requests.utils.default_headers()

headers.update(
    {
        'User-Agent': 'My User Agent 1.0',
    }
)

def next_chapter_in_text(tag):
    """
    Returns True if the text of the tag contains "Next Chapter".
    We need this because there's a bit of variation in the text of the
    <a href= ...> tag, so we can't just use find(string="Next Chapter").
    """
    if (tag.name == 'a' and
            tag.string != None and
            'next chapter' in tag.string.lower()):
        return True
    return False


def scrape_worm(start_link=FIRST_PAGE, maximum=None):
    """
    Scrapes every chapter, starting from the link given if provided.
    Returns:
        - A list of tuples: [(chapter_title, [list of words]), ...],
          (List and not dict because we want these in chapter order.)
        - The number of documents
        - The document frequencies as a dict (keys are the vocab).
    """
    link = start_link
    arcNum = 1;
    subNum = 1;
    chapter = 1;
    document = Document()

    # Keep scraping until the "Next Chapter" link is just "End"
    while(link != None):

        #grabs the next link to populate content.
        response = requests.get(link, headers=headers, timeout=10)
        content = bs(response.content, 'html.parser')
        # Find title of the chapter
        title = content.find(class_='entry-title').string

        if 'Interlude' not in title:
            arcList = re.findall(r"[-+]?\d*\.\d+|\d+", title)
            arcString = arcList[0] #get the arc and chapter as a string
            arcNum = math.floor(float(arcString)) #get the arc num as a num
            subList = arcString.split(".", 1)
            subNum = int(subList[1])
            # print(subNum)
            chapter = arcNum

            #we want to split up the book, with a document for every 5 arcs
            if(chapter > 1):
                if(chapter % 5 is 1):
                    #save old chap when we get to the first sub chapter of the new doc
                    if(subNum is 1):
                        print('Saving: Arc%d-%d.docx' % ((chapter - 5), (chapter - 1)))
                        document.save('Arc%d-%d.docx' % ((chapter - 5), (chapter - 1)))
                        document = Document()

            print(title)
            document.add_heading(title, level=1)
            # Find next link
            next_chapter_tag = content.find(next_chapter_in_text)
            if next_chapter_tag != None:
                link = next_chapter_tag.get('href')
                # link = None
            else:
                link = None

            # Chapter content
            entry_content = content.find(class_='entry-content')
            chapter_text_tags = entry_content.find_all('p')
            for paragraph_tag in chapter_text_tags:
                para_ref = document.add_paragraph()
                para_ref = keep_original_format(para_ref, paragraph_tag)

            document.add_page_break()
        else :
            print('skipping: ' + title)
            # Find next link
            next_chapter_tag = content.find(next_chapter_in_text)
            if next_chapter_tag != None:
                link = next_chapter_tag.get('href')
                # link = None
            else:
                link = None
    print('Saving: Arc-Last.docx')
    document.save('Arc-Last.docx')

def keep_original_format(ref, tag):
    try:
        isItalics = True
        text = str(tag) + "<em>" # guaruntee at least a single <em>. This is dumb but im tired.
        if('Next Chapter' in text or 'Last Chapter' in text):
            return ref
        results = text.split("<em>", 1)
        ref.add_run(clean(results[0]))
        while(results[1] != ''):
            if(isItalics):
                results = results[1].split("</em>", 1)
                ref.add_run(clean(results[0])).italic = True
                isItalics = not isItalics
            else:
                results = results[1].split("<em>", 1)
                ref.add_run(clean(results[0]))
                isItalics = not isItalics
        return ref
    except ValueError:
        print('error happened with this tag: ')
        print(tag)
        return ref

def clean(tag):
    text = TAG_RE.sub('', tag)
    return text

if __name__ == '__main__':
    # Start off in Gestation 1.1
    scrape_worm()
